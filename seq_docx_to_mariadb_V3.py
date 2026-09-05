#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
seq_docx_to_mariadb.py

DataStage SEQ 운영자 사전 질문지(개정본) DOCX를 읽어 MariaDB에 적재하고,
Airflow + dbt + Cosmos 전환을 위한 dbt tag 후보를 생성한다.

지원 문서:
  DataStage_SEQ_Airflow_DAG_운영자_사전질문지_개정본.docx

핵심 설계:
  1) 질문 항목명(좌측 셀)을 기준으로 파싱하므로 표 행 순서가 일부 바뀌어도 동작한다.
  2) Word Content Control(w:sdt) 체크박스의 실제 checked 상태까지 OOXML에서 읽는다.
  3) 새로 추가된 검증 기준/절차 항목을 구조화해 저장한다.
  4) 사람이 판단해야 하는 domain_* / run_* 은 임의 생성하지 않는다.
  5) cycle / trigger / operation / verification / master 계열은 명확한 답변만 규칙 기반으로 tag 후보를 만든다.
  6) 동일 SEQ 재적재 시 UPSERT하고 tag 후보는 재생성한다.

필수 패키지:
  pip install python-docx pymysql python-dotenv

환경변수(.env):
  MARIADB_HOST            (기본 127.0.0.1)
  MARIADB_PORT            (기본 3306)
  MARIADB_USER
  MARIADB_PASSWORD
  MARIADB_DATABASE
  MARIADB_CONNECT_TIMEOUT (기본 10초)

MariaDB Docker Compose 초기화용 환경변수(.env):
  MARIADB_ROOT_PASSWORD
  위 MARIADB_DATABASE / MARIADB_USER / MARIADB_PASSWORD를 동일하게 사용

사용 예:
  # 파싱 결과만 확인
  python seq_docx_to_mariadb.py ./answer.docx --dry-run

  # 단일 파일 적재
  python seq_docx_to_mariadb.py ./answer.docx

  # 디렉터리 내 DOCX 일괄 적재
  python seq_docx_to_mariadb.py ./answers --recursive

  # .env 파일을 명시적으로 지정
  python seq_docx_to_mariadb.py ./answers --env-file ./.env
"""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import sys
from dataclasses import dataclass, asdict
from datetime import date, datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError as exc:  # pragma: no cover
    raise SystemExit("python-docx가 필요합니다: pip install python-docx") from exc

try:
    import pymysql
except ImportError:  # dry-run에서는 DB 드라이버가 없어도 되도록 아래에서 체크
    pymysql = None

try:
    from dotenv import load_dotenv
except ImportError:  # dry-run은 .env 없이도 사용할 수 있도록 실행 시점에 체크
    load_dotenv = None


# -----------------------------------------------------------------------------
# 질문지 정의
# -----------------------------------------------------------------------------

@dataclass(frozen=True)
class FieldDef:
    label: str
    column: str
    kind: str = "text"  # text | choices | single
    options: Tuple[str, ...] = ()
    placeholder: str = ""


FIELDS: Tuple[FieldDef, ...] = (
    # 1. 기본 정보
    FieldDef("SEQ명", "seq_name"),
    FieldDef("업무명", "business_name"),
    FieldDef("운영 부서 / 담당자", "owner"),
    FieldDef("실행 주기", "execution_cycle", "single", ("일배치", "주배치", "월배치", "수시", "기타")),
    FieldDef("운영 중요도", "operational_importance", "single", ("상", "중", "하")),
    FieldDef("목표 완료시간", "target_completion_time", placeholder="예: 매일 07:30 이전"),
    FieldDef("주요 처리 목적", "main_processing_purpose", placeholder="예: 전일 점포 매출 집계 / 상품 기준정보 갱신"),
    FieldDef("처리 기준일 / 대상 기간", "processing_base_period", placeholder="예: 전일 / 당일 / 전월 / 월초~말일"),
    FieldDef("주요 결과물 / 사용처", "main_output_usage", placeholder="예: DW/DM 테이블, 보고서, 파일, 후속 업무"),
    FieldDef("주요 이용 업무 / 사용자", "main_user_business", placeholder="예: 점포 매출조회, 정산팀, BI 보고서"),

    # 2. 실행 흐름 및 선후행 관계
    FieldDef(
        "실행 시작 조건",
        "start_condition",
        "choices",
        ("정해진 시간", "선행 SEQ/Job 완료", "파일 도착", "데이터 생성", "수동 실행", "기타"),
    ),
    FieldDef("시작 조건 상세", "start_condition_detail", placeholder="대상/시간/조건을 입력해 주세요"),
    FieldDef("주요 처리 흐름", "main_processing_flow", placeholder="예: JOB_A → JOB_B → JOB_C"),
    FieldDef("병렬 실행 여부", "parallel_execution", "single", ("있음", "없음", "모름/확인 필요")),
    FieldDef("병렬/분기 상세", "parallel_branch_detail", placeholder="동시 실행 Job 또는 분기 조건을 입력해 주세요"),
    FieldDef("외부 의존성", "external_dependency", placeholder="다른 SEQ/Job, 파일, DB, 외부시스템 등을 입력해 주세요"),
    FieldDef(
        "함께 운영되는 SEQ/업무",
        "co_operated_seq_business",
        placeholder="현재 하나의 업무로 묶어 관리하거나 연속 실행하는 SEQ/업무",
    ),
    FieldDef(
        "후행 업무 / 사용처",
        "downstream_business_usage",
        placeholder="SEQ 완료 후 결과를 사용하는 업무 또는 후속 작업",
    ),
    FieldDef(
        "정상 완료 판단",
        "normal_completion_criteria",
        placeholder="예: SEQ 성공만 확인 / 결과 건수·금액·보고서까지 확인",
    ),

    # 3. 장애 및 재처리 방식
    FieldDef(
        "Job 실패 시 처리",
        "job_failure_action",
        "choices",
        ("SEQ 전체 종료", "다음 단계 계속", "오류 Job 실행", "운영자 확인 후 처리", "기타"),
    ),
    FieldDef("실패 처리 상세", "failure_action_detail", placeholder="실패 시 실제 운영 절차를 입력해 주세요"),
    FieldDef("자동 재시도", "auto_retry", "single", ("없음", "있음", "Job별 상이")),
    FieldDef("재시도 기준", "retry_criteria", placeholder="예: 3회 / 10분 간격 / 특정 오류만 재시도"),
    FieldDef(
        "재처리 방식",
        "reprocessing_method",
        "choices",
        ("전체 재실행", "실패 Job부터", "특정 구간부터", "상황별 판단"),
    ),
    FieldDef("재처리 상세", "reprocessing_detail", placeholder="재실행 시 선행 확인사항과 시작 지점을 입력해 주세요"),
    FieldDef("수동 운영 작업", "manual_operation", placeholder="수동 실행·중지·Skip·강제 종료·재기동 등을 입력해 주세요"),
    FieldDef(
        "재처리 전 확인사항",
        "pre_reprocessing_check",
        placeholder="예: 기존 데이터 삭제 여부, 선행 데이터 생성 여부, 기준일 확인",
    ),
    FieldDef(
        "재처리 후 확인방법",
        "post_reprocessing_check",
        placeholder="예: 건수/금액/보고서 확인, 운영자 육안 확인",
    ),
    FieldDef(
        "Skip/Reject 데이터 처리",
        "skip_reject_handling",
        placeholder="예: 별도 보관 / 수동 보정 / 무시 / 해당없음",
    ),

    # 4. 파라미터 / 알림 / 예외 규칙
    FieldDef("주요 파라미터", "main_parameters", placeholder="예: 기준일자, 경로, 시스템 구분, 재처리 여부"),
    FieldDef("파일/데이터 대기", "file_data_wait", placeholder="대기 대상, 확인 주기, 타임아웃 기준을 입력해 주세요"),
    FieldDef("성공/실패 알림", "success_failure_notification", placeholder="알림 조건, 방식, 대상자를 입력해 주세요"),
    FieldDef("예외 실행 규칙", "exception_execution_rule", placeholder="월말/휴일/특정 요일/특정 데이터 조건 등을 입력해 주세요"),
    FieldDef("정기 외 실행 사유", "non_regular_execution_reason", placeholder="예: 장애복구 / 데이터 정정 / 월말 추가 / 사용자 요청"),
    FieldDef("운영 확인 화면 / 로그", "operation_check_screen_log", placeholder="예: Director 로그, SQL 조회, 보고서, 파일 생성 결과"),

    # 5. 운영상 중요 확인사항
    FieldDef("가장 주의할 구간", "most_attention_section", placeholder="장애가 자주 발생하거나 운영 판단이 필요한 구간"),
    FieldDef("현재 운영 불편사항", "current_operation_pain_point", placeholder="수동 확인, 반복 작업, 재처리 어려움 등을 입력해 주세요"),
    FieldDef("전환 시 반드시 유지할 방식", "must_keep_after_transition", placeholder="Airflow 전환 후에도 유지되어야 하는 운영 규칙"),
    FieldDef("인터뷰 시 추가 확인", "interview_additional_check", placeholder="사전 작성이 어려워 인터뷰에서 확인할 내용을 입력해 주세요"),
    FieldDef("지연/실패 시 영향", "delay_failure_impact", placeholder="영향 받는 후속 업무, 사용자, 보고서, 마감 등을 입력해 주세요"),
    FieldDef("업무상 핵심 결과", "business_critical_result", placeholder="반드시 정상이어야 하는 주요 테이블/수치/파일/보고서"),

    # 6. 검증 기준 및 확인 절차
    FieldDef("검증 수행 시점", "verification_timing", placeholder="예: 매일 실행 후 / 마감 후 / 장애·재처리 시 / 월말"),
    FieldDef(
        "현재 결과 확인 방법",
        "verification_method",
        "choices",
        ("건수", "금액/수량 합계", "주요 Key/샘플", "보고서/화면", "파일", "기타"),
    ),
    FieldDef(
        "비교 기준 / 대상",
        "comparison_reference",
        "choices",
        ("원천 데이터", "기존 DW/DM", "전일/전월 결과", "보고서", "외부 파일/시스템", "기타"),
    ),
    FieldDef("주요 검증 대상", "main_verification_target", placeholder="예: 테이블명, 주요 금액/수량 컬럼, 핵심 보고서/파일"),
    FieldDef("NULL/중복/누락 확인", "null_duplicate_missing_check", placeholder="예: 필수값 NULL, Key 중복, 특정 일자/점포 데이터 누락 여부"),
    FieldDef(
        "결과 차이 허용 여부",
        "difference_tolerance",
        "single",
        ("완전 일치 필요", "일부 차이 허용", "상황별 판단", "모름/확인 필요"),
    ),
    FieldDef(
        "허용 차이 / 정상 예외",
        "allowed_difference_normal_exception",
        placeholder="예: 마스터 변경, 이력 기간 차이, 지연 수신 데이터 등 정상적으로 차이가 날 수 있는 경우",
    ),
    FieldDef("MASTER/기준정보 영향", "master_reference", "single", ("있음", "없음", "모름/확인 필요")),
    FieldDef("과거 MASTER 시점 확인", "historical_master_available", "single", ("가능", "불가능", "일부 가능", "모름/확인 필요")),
    FieldDef("이력 데이터 비교 가능 기간", "history_comparable_period", placeholder="예: 최근 4개월 / 1년 / 특정 기준일 이후 / 확인 필요"),
    FieldDef(
        "과거일자 재실행 결과",
        "past_date_rerun_result",
        "single",
        ("동일 결과 예상", "MASTER/이력 시점에 따라 달라질 수 있음", "모름/확인 필요"),
    ),
    FieldDef("검증 실패 시 처리", "verification_failure_action", placeholder="예: 재실행 / 데이터 보정 / 원인 확인 후 승인 / 담당자 확인"),
)

FIELD_BY_LABEL: Dict[str, FieldDef] = {f.label: f for f in FIELDS}

MATERIAL_OPTIONS: Tuple[str, ...] = (
    "DataStage SEQ 캡처",
    "SEQ Export(DSX/XML)",
    "운영 매뉴얼",
    "장애/재처리 이력",
    "스케줄/선후행 목록",
    "기타 자료",
)

# python-docx cell.text가 Content Control 안의 텍스트를 누락할 수 있으므로
# 아래 OOXML namespace를 직접 사용한다.
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"


# -----------------------------------------------------------------------------
# DOCX XML 처리
# -----------------------------------------------------------------------------

def _local_name(tag: str) -> str:
    return tag.rsplit("}", 1)[-1] if "}" in tag else tag


def _is_checkbox_sdt(element: Any) -> bool:
    if _local_name(element.tag) != "sdt":
        return False
    for node in element.iter():
        if node.tag == f"{{{W14_NS}}}checkbox":
            return True
    return False


def _checkbox_checked(element: Any) -> bool:
    for node in element.iter():
        if node.tag == f"{{{W14_NS}}}checked":
            val = node.get(f"{{{W14_NS}}}val") or node.get("val") or "0"
            return str(val).lower() in {"1", "true", "on"}
    return False


def _extract_xml_text(element: Any) -> str:
    """OOXML 노드를 문서 순서대로 읽어 Content Control까지 포함한 표시 텍스트를 만든다."""
    name = _local_name(element.tag)

    if name == "sdt" and _is_checkbox_sdt(element):
        return "☒" if _checkbox_checked(element) else "☐"

    if name == "t":
        return element.text or ""
    if name == "tab":
        return "\t"
    if name in {"br", "cr"}:
        return "\n"

    parts: List[str] = []
    for child in element:
        parts.append(_extract_xml_text(child))
    return "".join(parts)


def extract_cell_text(cell: Any) -> str:
    """Content Control/체크박스를 포함해 셀의 실제 내용을 반환한다."""
    text = _extract_xml_text(cell._tc)
    return normalize_whitespace(text, keep_newline=True)


def extract_paragraph_text(paragraph: Any) -> str:
    return normalize_whitespace(_extract_xml_text(paragraph._p), keep_newline=False)


def normalize_whitespace(value: str, keep_newline: bool = False) -> str:
    value = value.replace("\xa0", " ").replace("\u200b", "")
    if keep_newline:
        lines = [re.sub(r"[\t ]+", " ", line).strip() for line in value.splitlines()]
        return "\n".join(line for line in lines if line)
    return re.sub(r"\s+", " ", value).strip()


def clean_label(value: str) -> str:
    return normalize_whitespace(value).replace("Job ", "Job ")


def clean_text_answer(raw: str, placeholder: str = "") -> str:
    text = normalize_whitespace(raw, keep_newline=True)
    if not text:
        return ""
    if placeholder and normalize_whitespace(text) == normalize_whitespace(placeholder):
        return ""
    # Word Content Control 기본 placeholder
    if normalize_whitespace(text) in {"클릭하여 입력", "YYYY-MM-DD"}:
        return ""
    return text


# -----------------------------------------------------------------------------
# 체크박스/선택값 파싱
# -----------------------------------------------------------------------------

SELECTED_MARKS = ("☒", "☑", "✅", "✔", "✓", "■")
UNCHECKED_MARKS = ("☐", "□")


def parse_choices(rendered_text: str, options: Sequence[str]) -> List[str]:
    """
    선택형 셀에서 체크된 옵션을 반환한다.

    - Content Control checkbox는 extract_cell_text가 ☒/☐ 로 렌더링한다.
    - 일반 Unicode checkbox도 지원한다.
    - 사용자가 체크박스를 지우고 옵션 하나만 남긴 경우도 지원한다.
    """
    text = normalize_whitespace(rendered_text)
    selected: List[str] = []

    for opt in options:
        escaped = re.escape(opt)
        mark_pattern = "|".join(re.escape(m) for m in SELECTED_MARKS)
        if re.search(rf"(?:{mark_pattern})\s*{escaped}(?=\s|$)", text, flags=re.IGNORECASE):
            selected.append(opt)
            continue
        # [x] / (x) / x 형태
        if re.search(rf"(?:\[\s*[xXvV]\s*\]|\(\s*[xXvV]\s*\))\s*{escaped}(?=\s|$)", text):
            selected.append(opt)

    if selected:
        return selected

    # 체크 기호가 하나도 없고 전체 텍스트가 특정 옵션 하나와 동일한 경우
    # (운영자가 선택하지 않은 옵션을 삭제해서 작성한 문서 지원)
    stripped = text.strip(" ,;/|")
    exact = [opt for opt in options if stripped == opt]
    if len(exact) == 1:
        return exact

    return []


def choice_code(selected: Sequence[str]) -> str:
    return selected[0] if len(selected) == 1 else ""


# -----------------------------------------------------------------------------
# 문서 파싱
# -----------------------------------------------------------------------------

@dataclass
class ParsedQuestionnaire:
    source_file: str
    source_path: str
    source_sha256: str
    answers: Dict[str, str]
    selected: Dict[str, List[str]]
    available_materials: List[str]
    completion_date: Optional[str]
    author: str
    unmapped_labels: Dict[str, str]
    tags: List[Dict[str, Any]]


def file_sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def parse_footer_metadata(doc: Any) -> Tuple[Optional[str], str]:
    completion_date: Optional[str] = None
    author = ""

    for p in doc.paragraphs:
        text = extract_paragraph_text(p)
        if "작성 완료일" not in text or "작성자" not in text:
            continue

        # Content Control 내용까지 포함한 예: 작성 완료일: 2026-09-05 작성자: 홍길동
        m = re.search(
            r"작성\s*완료일\s*:\s*(.*?)\s+작성자\s*:\s*(.*)$",
            text,
        )
        if not m:
            continue
        date_text = clean_text_answer(m.group(1), "YYYY-MM-DD")
        author = clean_text_answer(m.group(2), "클릭하여 입력")
        completion_date = normalize_date_string(date_text)
        break

    return completion_date, author


def normalize_date_string(value: str) -> Optional[str]:
    if not value:
        return None
    value = value.strip()
    candidates = ["%Y-%m-%d", "%Y.%m.%d", "%Y/%m/%d", "%Y%m%d"]
    for fmt in candidates:
        try:
            return datetime.strptime(value, fmt).date().isoformat()
        except ValueError:
            pass
    # 날짜가 아닌 자유기술은 DB DATE에 넣지 않고 raw payload에만 남긴다.
    return None


def parse_questionnaire(path: Path) -> ParsedQuestionnaire:
    doc = Document(str(path))
    answers: Dict[str, str] = {f.column: "" for f in FIELDS}
    selected: Dict[str, List[str]] = {}
    unmapped_labels: Dict[str, str] = {}

    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            label = clean_label(extract_cell_text(row.cells[0]))
            if not label:
                continue
            value = extract_cell_text(row.cells[1])
            field = FIELD_BY_LABEL.get(label)
            if field is None:
                # 질문처럼 보이는 2열 행만 기록한다.
                if value:
                    unmapped_labels[label] = value
                continue

            if field.kind in {"choices", "single"}:
                answers[field.column] = normalize_whitespace(value)
                selected[field.column] = parse_choices(value, field.options)
            else:
                answers[field.column] = clean_text_answer(value, field.placeholder)

    # 7. 사전 제공 가능 자료: 개정본은 1열 표
    available_materials: List[str] = []
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) != 1:
                continue
            value = extract_cell_text(row.cells[0])
            found = parse_choices(value, MATERIAL_OPTIONS)
            if found:
                available_materials = found

    completion_date, author = parse_footer_metadata(doc)

    parsed = ParsedQuestionnaire(
        source_file=path.name,
        source_path=str(path.resolve()),
        source_sha256=file_sha256(path),
        answers=answers,
        selected=selected,
        available_materials=available_materials,
        completion_date=completion_date,
        author=author,
        unmapped_labels=unmapped_labels,
        tags=[],
    )
    parsed.tags = generate_tag_candidates(parsed)
    return parsed


# -----------------------------------------------------------------------------
# Tag 후보 생성
# -----------------------------------------------------------------------------

NEGATIVE_TEXT = {
    "없음", "해당없음", "해당 없음", "없다", "n", "no", "아니오", "모름", "확인 필요", "모름/확인 필요"
}


def has_positive_text(value: str) -> bool:
    v = normalize_whitespace(value).lower()
    return bool(v) and v not in NEGATIVE_TEXT


def _tag(
    category: str,
    name: str,
    source_field: str,
    source_value: str,
    rule_code: str,
    confidence: float = 1.0,
    requires_review: bool = False,
) -> Dict[str, Any]:
    return {
        "category": category,
        "tag_name": name,
        "source_field": source_field,
        "source_value": source_value,
        "rule_code": rule_code,
        "confidence": confidence,
        "requires_review": 1 if requires_review else 0,
    }


def generate_tag_candidates(parsed: ParsedQuestionnaire) -> List[Dict[str, Any]]:
    a = parsed.answers
    s = parsed.selected
    tags: List[Dict[str, Any]] = []

    def add(category: str, name: str, field: str, value: str, rule: str, confidence: float = 1.0, review: bool = False) -> None:
        tags.append(_tag(category, name, field, value, rule, confidence, review))

    # 실행 주기
    cycle_map = {
        "일배치": "cycle_daily",
        "주배치": "cycle_weekly",
        "월배치": "cycle_monthly",
        "수시": "cycle_adhoc",
        "기타": "cycle_other",
    }
    for value in s.get("execution_cycle", []):
        if value in cycle_map:
            add("cycle", cycle_map[value], "execution_cycle", value, "CYCLE_001")

    # 운영 중요도: '상'만 critical 후보로 생성. 중/하는 tag보다 meta 성격이 강하므로 raw 값으로 보존.
    if "상" in s.get("operational_importance", []):
        add("operation", "critical", "operational_importance", "상", "OPS_001")

    # 시작 조건
    start_map = {
        "정해진 시간": "trigger_schedule",
        "선행 SEQ/Job 완료": "trigger_dependency",
        "파일 도착": "trigger_file",
        "데이터 생성": "trigger_data",
        "수동 실행": "run_manual",
        "기타": "trigger_other",
    }
    for value in s.get("start_condition", []):
        if value in start_map:
            category = "run" if value == "수동 실행" else "trigger"
            add(category, start_map[value], "start_condition", value, "TRG_001")

    # 외부 의존/대기
    if has_positive_text(a.get("external_dependency", "")) or has_positive_text(a.get("file_data_wait", "")):
        src = a.get("external_dependency", "") or a.get("file_data_wait", "")
        add("trigger", "external_dependency", "external_dependency", src, "TRG_002", 0.95)

    # 병렬성
    if "있음" in s.get("parallel_execution", []):
        add("operation", "parallel_run", "parallel_execution", "있음", "OPS_002")

    # 재처리
    rerun_map = {
        "전체 재실행": "reprocess_full",
        "실패 Job부터": "reprocess_failed",
        "특정 구간부터": "reprocess_partial",
        "상황별 판단": "reprocess_conditional",
    }
    rerun_values = s.get("reprocessing_method", [])
    if rerun_values or has_positive_text(a.get("reprocessing_detail", "")):
        add("operation", "reprocessable", "reprocessing_method", ", ".join(rerun_values) or a.get("reprocessing_detail", ""), "RERUN_001")
    for value in rerun_values:
        if value in rerun_map:
            add("operation", rerun_map[value], "reprocessing_method", value, "RERUN_002")

    # 자동 재시도
    if "있음" in s.get("auto_retry", []) or "Job별 상이" in s.get("auto_retry", []):
        value = ", ".join(s.get("auto_retry", []))
        add("operation", "retry_available", "auto_retry", value, "OPS_003")

    # 수동 운영 / 파라미터 / 예외 / 알림
    if has_positive_text(a.get("manual_operation", "")):
        add("operation", "manual_operation", "manual_operation", a["manual_operation"], "OPS_004")
    if has_positive_text(a.get("main_parameters", "")):
        add("operation", "parameterized", "main_parameters", a["main_parameters"], "OPS_005")
    if has_positive_text(a.get("exception_execution_rule", "")):
        add("operation", "special_schedule", "exception_execution_rule", a["exception_execution_rule"], "OPS_006", 0.9)
    if has_positive_text(a.get("non_regular_execution_reason", "")):
        add("run", "run_adhoc_possible", "non_regular_execution_reason", a["non_regular_execution_reason"], "OPS_007", 0.9)
    if has_positive_text(a.get("success_failure_notification", "")):
        add("operation", "alert_required", "success_failure_notification", a["success_failure_notification"], "OPS_008", 0.9)

    # 검증 방식
    verify_map = {
        "건수": "verify_count",
        "금액/수량 합계": "verify_sum",
        "주요 Key/샘플": "verify_sample",
        "보고서/화면": "verify_report",
        "파일": "verify_file",
        "기타": "verify_other",
    }
    for value in s.get("verification_method", []):
        if value in verify_map:
            add("verification", verify_map[value], "verification_method", value, "VER_001")

    # 비교 기준은 실행 tag라기보다는 검증 policy 메타데이터에 가깝지만 후보 tag로 남긴다.
    ref_map = {
        "원천 데이터": "verify_ref_source",
        "기존 DW/DM": "verify_ref_legacy_dw_dm",
        "전일/전월 결과": "verify_ref_prior_period",
        "보고서": "verify_ref_report",
        "외부 파일/시스템": "verify_ref_external",
        "기타": "verify_ref_other",
    }
    for value in s.get("comparison_reference", []):
        if value in ref_map:
            add("verification_reference", ref_map[value], "comparison_reference", value, "VER_002")

    # 검증 정책
    tolerance_map = {
        "완전 일치 필요": "policy_strict",
        "일부 차이 허용": "policy_tolerance",
        "상황별 판단": "policy_conditional",
        "모름/확인 필요": "policy_check_needed",
    }
    for value in s.get("difference_tolerance", []):
        if value in tolerance_map:
            add("verification_policy", tolerance_map[value], "difference_tolerance", value, "VER_003")

    # NULL/중복/누락을 실제 운영에서 확인한다고 기술한 경우
    if has_positive_text(a.get("null_duplicate_missing_check", "")):
        add("verification", "verify_data_quality", "null_duplicate_missing_check", a["null_duplicate_missing_check"], "VER_004", 0.9)

    # MASTER / 시점
    if "있음" in s.get("master_reference", []):
        add("data", "data_master_ref", "master_reference", "있음", "DATA_001")
    elif "모름/확인 필요" in s.get("master_reference", []):
        add("data", "data_master_check_needed", "master_reference", "모름/확인 필요", "DATA_002")

    master_time_map = {
        "가능": "master_time_known",
        "불가능": "master_time_unknown",
        "일부 가능": "master_time_partial",
        "모름/확인 필요": "master_time_check_needed",
    }
    for value in s.get("historical_master_available", []):
        if value in master_time_map:
            add("data", master_time_map[value], "historical_master_available", value, "DATA_003")

    # 이력 비교기간이 구체적으로 작성된 경우에만 history 후보 생성
    history_period = a.get("history_comparable_period", "")
    if has_positive_text(history_period) and "확인 필요" not in history_period:
        add("data", "data_history", "history_comparable_period", history_period, "DATA_004", 0.85, True)

    rerun_result_map = {
        "동일 결과 예상": "rerun_deterministic",
        "MASTER/이력 시점에 따라 달라질 수 있음": "rerun_time_sensitive",
        "모름/확인 필요": "rerun_result_check_needed",
    }
    for value in s.get("past_date_rerun_result", []):
        if value in rerun_result_map:
            add("verification_policy", rerun_result_map[value], "past_date_rerun_result", value, "VER_005")

    # 중복 tag 제거: 동일 이름은 최초 규칙 유지
    dedup: Dict[str, Dict[str, Any]] = {}
    for item in tags:
        dedup.setdefault(item["tag_name"], item)
    return list(dedup.values())


# -----------------------------------------------------------------------------
# MariaDB DDL / 적재
# -----------------------------------------------------------------------------

MAIN_TABLE = "seq_operator_questionnaire"
TAG_TABLE = "seq_operator_questionnaire_tag"

DDL_MAIN = f"""
CREATE TABLE IF NOT EXISTS {MAIN_TABLE} (
    id BIGINT NOT NULL AUTO_INCREMENT,
    seq_name VARCHAR(255) NOT NULL,
    business_name VARCHAR(255) NULL,
    owner VARCHAR(500) NULL,

    execution_cycle_raw TEXT NULL,
    execution_cycle_json LONGTEXT NULL,
    operational_importance_raw TEXT NULL,
    operational_importance_code VARCHAR(100) NULL,
    target_completion_time VARCHAR(255) NULL,
    main_processing_purpose TEXT NULL,
    processing_base_period TEXT NULL,
    main_output_usage TEXT NULL,
    main_user_business TEXT NULL,

    start_condition_raw TEXT NULL,
    start_condition_json LONGTEXT NULL,
    start_condition_detail TEXT NULL,
    main_processing_flow LONGTEXT NULL,
    parallel_execution_raw TEXT NULL,
    parallel_execution_code VARCHAR(100) NULL,
    parallel_branch_detail LONGTEXT NULL,
    external_dependency LONGTEXT NULL,
    co_operated_seq_business LONGTEXT NULL,
    downstream_business_usage LONGTEXT NULL,
    normal_completion_criteria LONGTEXT NULL,

    job_failure_action_raw TEXT NULL,
    job_failure_action_json LONGTEXT NULL,
    failure_action_detail LONGTEXT NULL,
    auto_retry_raw TEXT NULL,
    auto_retry_code VARCHAR(100) NULL,
    retry_criteria TEXT NULL,
    reprocessing_method_raw TEXT NULL,
    reprocessing_method_json LONGTEXT NULL,
    reprocessing_detail LONGTEXT NULL,
    manual_operation LONGTEXT NULL,
    pre_reprocessing_check LONGTEXT NULL,
    post_reprocessing_check LONGTEXT NULL,
    skip_reject_handling LONGTEXT NULL,

    main_parameters LONGTEXT NULL,
    file_data_wait LONGTEXT NULL,
    success_failure_notification LONGTEXT NULL,
    exception_execution_rule LONGTEXT NULL,
    non_regular_execution_reason LONGTEXT NULL,
    operation_check_screen_log LONGTEXT NULL,

    most_attention_section LONGTEXT NULL,
    current_operation_pain_point LONGTEXT NULL,
    must_keep_after_transition LONGTEXT NULL,
    interview_additional_check LONGTEXT NULL,
    delay_failure_impact LONGTEXT NULL,
    business_critical_result LONGTEXT NULL,

    verification_timing LONGTEXT NULL,
    verification_method_raw TEXT NULL,
    verification_method_json LONGTEXT NULL,
    comparison_reference_raw TEXT NULL,
    comparison_reference_json LONGTEXT NULL,
    main_verification_target LONGTEXT NULL,
    null_duplicate_missing_check LONGTEXT NULL,
    difference_tolerance_raw TEXT NULL,
    difference_tolerance_code VARCHAR(100) NULL,
    allowed_difference_normal_exception LONGTEXT NULL,
    master_reference_raw TEXT NULL,
    master_reference_code VARCHAR(100) NULL,
    historical_master_available_raw TEXT NULL,
    historical_master_available_code VARCHAR(100) NULL,
    history_comparable_period TEXT NULL,
    past_date_rerun_result_raw TEXT NULL,
    past_date_rerun_result_code VARCHAR(255) NULL,
    verification_failure_action LONGTEXT NULL,

    available_materials_json LONGTEXT NULL,
    completed_date DATE NULL,
    author VARCHAR(255) NULL,

    derived_tags_json LONGTEXT NULL,
    raw_answers_json LONGTEXT NULL,
    unmapped_labels_json LONGTEXT NULL,

    source_file VARCHAR(500) NOT NULL,
    source_path VARCHAR(1500) NULL,
    source_sha256 CHAR(64) NOT NULL,
    imported_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,
    updated_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,

    PRIMARY KEY (id),
    UNIQUE KEY uk_seq_operator_questionnaire_seq (seq_name),
    KEY idx_seq_operator_questionnaire_business (business_name),
    KEY idx_seq_operator_questionnaire_sha (source_sha256)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;
"""

DDL_TAG = f"""
CREATE TABLE IF NOT EXISTS {TAG_TABLE} (
    id BIGINT NOT NULL AUTO_INCREMENT,
    questionnaire_id BIGINT NOT NULL,
    tag_category VARCHAR(64) NOT NULL,
    tag_name VARCHAR(128) NOT NULL,
    source_field VARCHAR(128) NULL,
    source_value TEXT NULL,
    rule_code VARCHAR(64) NULL,
    confidence DECIMAL(5,4) NOT NULL DEFAULT 1.0000,
    requires_review TINYINT(1) NOT NULL DEFAULT 0,
    created_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,

    PRIMARY KEY (id),
    UNIQUE KEY uk_seq_questionnaire_tag (questionnaire_id, tag_name),
    KEY idx_seq_questionnaire_tag_name (tag_name),
    KEY idx_seq_questionnaire_tag_category (tag_category),
    CONSTRAINT fk_seq_questionnaire_tag
        FOREIGN KEY (questionnaire_id) REFERENCES {MAIN_TABLE}(id)
        ON DELETE CASCADE
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;
"""


def jdump(value: Any) -> str:
    return json.dumps(value, ensure_ascii=False, separators=(",", ":"))


def parsed_to_row(parsed: ParsedQuestionnaire) -> Dict[str, Any]:
    a = parsed.answers
    s = parsed.selected

    row: Dict[str, Any] = {
        "seq_name": a.get("seq_name", ""),
        "business_name": a.get("business_name", ""),
        "owner": a.get("owner", ""),

        "execution_cycle_raw": a.get("execution_cycle", ""),
        "execution_cycle_json": jdump(s.get("execution_cycle", [])),
        "operational_importance_raw": a.get("operational_importance", ""),
        "operational_importance_code": choice_code(s.get("operational_importance", [])),
        "target_completion_time": a.get("target_completion_time", ""),
        "main_processing_purpose": a.get("main_processing_purpose", ""),
        "processing_base_period": a.get("processing_base_period", ""),
        "main_output_usage": a.get("main_output_usage", ""),
        "main_user_business": a.get("main_user_business", ""),

        "start_condition_raw": a.get("start_condition", ""),
        "start_condition_json": jdump(s.get("start_condition", [])),
        "start_condition_detail": a.get("start_condition_detail", ""),
        "main_processing_flow": a.get("main_processing_flow", ""),
        "parallel_execution_raw": a.get("parallel_execution", ""),
        "parallel_execution_code": choice_code(s.get("parallel_execution", [])),
        "parallel_branch_detail": a.get("parallel_branch_detail", ""),
        "external_dependency": a.get("external_dependency", ""),
        "co_operated_seq_business": a.get("co_operated_seq_business", ""),
        "downstream_business_usage": a.get("downstream_business_usage", ""),
        "normal_completion_criteria": a.get("normal_completion_criteria", ""),

        "job_failure_action_raw": a.get("job_failure_action", ""),
        "job_failure_action_json": jdump(s.get("job_failure_action", [])),
        "failure_action_detail": a.get("failure_action_detail", ""),
        "auto_retry_raw": a.get("auto_retry", ""),
        "auto_retry_code": choice_code(s.get("auto_retry", [])),
        "retry_criteria": a.get("retry_criteria", ""),
        "reprocessing_method_raw": a.get("reprocessing_method", ""),
        "reprocessing_method_json": jdump(s.get("reprocessing_method", [])),
        "reprocessing_detail": a.get("reprocessing_detail", ""),
        "manual_operation": a.get("manual_operation", ""),
        "pre_reprocessing_check": a.get("pre_reprocessing_check", ""),
        "post_reprocessing_check": a.get("post_reprocessing_check", ""),
        "skip_reject_handling": a.get("skip_reject_handling", ""),

        "main_parameters": a.get("main_parameters", ""),
        "file_data_wait": a.get("file_data_wait", ""),
        "success_failure_notification": a.get("success_failure_notification", ""),
        "exception_execution_rule": a.get("exception_execution_rule", ""),
        "non_regular_execution_reason": a.get("non_regular_execution_reason", ""),
        "operation_check_screen_log": a.get("operation_check_screen_log", ""),

        "most_attention_section": a.get("most_attention_section", ""),
        "current_operation_pain_point": a.get("current_operation_pain_point", ""),
        "must_keep_after_transition": a.get("must_keep_after_transition", ""),
        "interview_additional_check": a.get("interview_additional_check", ""),
        "delay_failure_impact": a.get("delay_failure_impact", ""),
        "business_critical_result": a.get("business_critical_result", ""),

        "verification_timing": a.get("verification_timing", ""),
        "verification_method_raw": a.get("verification_method", ""),
        "verification_method_json": jdump(s.get("verification_method", [])),
        "comparison_reference_raw": a.get("comparison_reference", ""),
        "comparison_reference_json": jdump(s.get("comparison_reference", [])),
        "main_verification_target": a.get("main_verification_target", ""),
        "null_duplicate_missing_check": a.get("null_duplicate_missing_check", ""),
        "difference_tolerance_raw": a.get("difference_tolerance", ""),
        "difference_tolerance_code": choice_code(s.get("difference_tolerance", [])),
        "allowed_difference_normal_exception": a.get("allowed_difference_normal_exception", ""),
        "master_reference_raw": a.get("master_reference", ""),
        "master_reference_code": choice_code(s.get("master_reference", [])),
        "historical_master_available_raw": a.get("historical_master_available", ""),
        "historical_master_available_code": choice_code(s.get("historical_master_available", [])),
        "history_comparable_period": a.get("history_comparable_period", ""),
        "past_date_rerun_result_raw": a.get("past_date_rerun_result", ""),
        "past_date_rerun_result_code": choice_code(s.get("past_date_rerun_result", [])),
        "verification_failure_action": a.get("verification_failure_action", ""),

        "available_materials_json": jdump(parsed.available_materials),
        "completed_date": parsed.completion_date,
        "author": parsed.author,

        "derived_tags_json": jdump([t["tag_name"] for t in parsed.tags]),
        "raw_answers_json": jdump({"answers": a, "selected": s}),
        "unmapped_labels_json": jdump(parsed.unmapped_labels),

        "source_file": parsed.source_file,
        "source_path": parsed.source_path,
        "source_sha256": parsed.source_sha256,
    }
    return row


def env_required(name: str) -> str:
    value = os.getenv(name, "").strip()
    if not value:
        raise RuntimeError(f"필수 환경변수가 없습니다: {name}")
    return value


def connect_db():
    """MariaDB 접속정보는 CLI가 아니라 현재 프로세스 환경변수에서만 읽는다."""
    if pymysql is None:
        raise RuntimeError("pymysql이 필요합니다: pip install pymysql")

    host = os.getenv("MARIADB_HOST", "127.0.0.1").strip() or "127.0.0.1"
    port = int(os.getenv("MARIADB_PORT", "3306"))
    user = env_required("MARIADB_USER")
    password = env_required("MARIADB_PASSWORD")
    database = env_required("MARIADB_DATABASE")
    connect_timeout = int(os.getenv("MARIADB_CONNECT_TIMEOUT", "10"))

    return pymysql.connect(
        host=host,
        port=port,
        user=user,
        password=password,
        database=database,
        charset="utf8mb4",
        autocommit=False,
        connect_timeout=connect_timeout,
        cursorclass=pymysql.cursors.DictCursor,
    )


def ensure_schema(conn: Any) -> None:
    with conn.cursor() as cur:
        cur.execute(DDL_MAIN)
        cur.execute(DDL_TAG)
    conn.commit()


def upsert_questionnaire(conn: Any, parsed: ParsedQuestionnaire) -> Tuple[int, str]:
    row = parsed_to_row(parsed)
    columns = list(row.keys())
    placeholders = ", ".join(["%s"] * len(columns))
    col_sql = ", ".join(f"`{c}`" for c in columns)
    update_cols = [c for c in columns if c != "seq_name"]
    update_sql = ", ".join(f"`{c}`=VALUES(`{c}`)" for c in update_cols)

    # LAST_INSERT_ID(id)를 사용하면 INSERT/UPDATE 모두 cursor.lastrowid로 id 획득 가능
    sql = f"""
        INSERT INTO {MAIN_TABLE} ({col_sql})
        VALUES ({placeholders})
        ON DUPLICATE KEY UPDATE
            {update_sql},
            id=LAST_INSERT_ID(id)
    """

    with conn.cursor() as cur:
        cur.execute(sql, [row[c] if row[c] != "" else None for c in columns])
        questionnaire_id = int(cur.lastrowid)

        # tag 후보는 현재 질문지 기준으로 완전 재생성
        cur.execute(f"DELETE FROM {TAG_TABLE} WHERE questionnaire_id=%s", (questionnaire_id,))
        if parsed.tags:
            tag_sql = f"""
                INSERT INTO {TAG_TABLE}
                    (questionnaire_id, tag_category, tag_name, source_field, source_value,
                     rule_code, confidence, requires_review)
                VALUES (%s,%s,%s,%s,%s,%s,%s,%s)
            """
            cur.executemany(
                tag_sql,
                [
                    (
                        questionnaire_id,
                        t["category"],
                        t["tag_name"],
                        t["source_field"],
                        t["source_value"],
                        t["rule_code"],
                        t["confidence"],
                        t["requires_review"],
                    )
                    for t in parsed.tags
                ],
            )
    conn.commit()
    return questionnaire_id, "UPSERT"


# -----------------------------------------------------------------------------
# CLI
# -----------------------------------------------------------------------------

def iter_docx_inputs(inputs: Sequence[str], recursive: bool) -> Iterable[Path]:
    seen: set[Path] = set()
    for item in inputs:
        path = Path(item).expanduser()
        if path.is_file():
            if path.suffix.lower() == ".docx" and path.name[0:2] != "~$":
                rp = path.resolve()
                if rp not in seen:
                    seen.add(rp)
                    yield path
            continue
        if path.is_dir():
            pattern = "**/*.docx" if recursive else "*.docx"
            for p in sorted(path.glob(pattern)):
                if p.name.startswith("~$"):
                    continue
                rp = p.resolve()
                if rp not in seen:
                    seen.add(rp)
                    yield p
            continue
        # shell glob을 따옴표로 넘긴 경우 간단 지원
        parent = path.parent if str(path.parent) != "" else Path(".")
        for p in sorted(parent.glob(path.name)):
            if p.is_file() and p.suffix.lower() == ".docx" and not p.name.startswith("~$"):
                rp = p.resolve()
                if rp not in seen:
                    seen.add(rp)
                    yield p


def build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(
        description="DataStage SEQ 운영자 사전질문지 개정본 DOCX → MariaDB 적재 + dbt tag 후보 생성"
    )
    p.add_argument("inputs", nargs="+", help="DOCX 파일, 디렉터리 또는 glob")
    p.add_argument("--recursive", action="store_true", help="디렉터리 하위까지 DOCX 검색")
    p.add_argument("--dry-run", action="store_true", help="DB 적재 없이 파싱 결과 JSON 출력")
    p.add_argument(
        "--allow-missing-seq",
        action="store_true",
        help="SEQ명이 비어 있어도 파일명으로 임시 SEQ명을 만들어 적재",
    )

    p.add_argument(
        "--env-file",
        default=os.getenv("ENV_FILE", ".env"),
        help="MariaDB 접속 환경변수를 읽을 .env 파일 경로 (기본: ./.env)",
    )
    return p


def print_dry_run(parsed: ParsedQuestionnaire) -> None:
    obj = {
        "source_file": parsed.source_file,
        "seq_name": parsed.answers.get("seq_name", ""),
        "business_name": parsed.answers.get("business_name", ""),
        "answers": parsed.answers,
        "selected": parsed.selected,
        "available_materials": parsed.available_materials,
        "completion_date": parsed.completion_date,
        "author": parsed.author,
        "tags": parsed.tags,
        "unmapped_labels": parsed.unmapped_labels,
    }
    print(json.dumps(obj, ensure_ascii=False, indent=2))


def load_environment(env_file: str, required: bool) -> None:
    """
    .env를 프로세스 환경으로 로드한다.

    우선순위는 OS/컨테이너에 이미 주입된 환경변수 > .env 이다.
    따라서 Docker/CI에서 환경변수를 직접 주입해도 동일 코드로 동작한다.
    """
    path = Path(env_file).expanduser() if env_file else None

    if path and path.exists():
        if load_dotenv is None:
            raise RuntimeError("python-dotenv가 필요합니다: pip install python-dotenv")
        load_dotenv(dotenv_path=path, override=False)
    elif required:
        # .env 파일이 없어도 Docker/CI에서 필수 변수가 이미 주입된 경우는 허용한다.
        needed = ("MARIADB_USER", "MARIADB_PASSWORD", "MARIADB_DATABASE")
        if not all(os.getenv(k, "").strip() for k in needed):
            raise RuntimeError(
                f".env 파일을 찾을 수 없습니다: {path}. "
                "또는 MARIADB_USER/MARIADB_PASSWORD/MARIADB_DATABASE를 환경변수로 주입하세요."
            )


def validate_db_env() -> None:
    missing = [
        name
        for name in ("MARIADB_USER", "MARIADB_PASSWORD", "MARIADB_DATABASE")
        if not os.getenv(name, "").strip()
    ]
    if missing:
        raise SystemExit("DB 환경변수가 부족합니다: " + ", ".join(missing))


def masked_db_target() -> str:
    host = os.getenv("MARIADB_HOST", "127.0.0.1")
    port = os.getenv("MARIADB_PORT", "3306")
    database = os.getenv("MARIADB_DATABASE", "")
    user = os.getenv("MARIADB_USER", "")
    return f"{user}@{host}:{port}/{database}"


def main() -> int:
    args = build_parser().parse_args()

    # dry-run은 DB가 필요 없지만, .env가 존재하면 동일하게 로드해 실행환경을 일관되게 유지한다.
    load_environment(args.env_file, required=not args.dry_run)

    files = list(iter_docx_inputs(args.inputs, args.recursive))
    if not files:
        print("처리할 DOCX 파일이 없습니다.", file=sys.stderr)
        return 2

    parsed_docs: List[ParsedQuestionnaire] = []
    errors: List[str] = []

    for path in files:
        try:
            parsed = parse_questionnaire(path)
            seq_name = parsed.answers.get("seq_name", "").strip()
            if not seq_name:
                if args.dry_run:
                    pass
                elif args.allow_missing_seq:
                    parsed.answers["seq_name"] = f"__FILE__:{path.stem}"
                else:
                    raise ValueError("SEQ명이 비어 있습니다. --allow-missing-seq로 임시 적재 가능")
            parsed_docs.append(parsed)
        except Exception as exc:
            errors.append(f"{path}: {exc}")

    if args.dry_run:
        for i, parsed in enumerate(parsed_docs):
            if i:
                print("\n" + "=" * 100 + "\n")
            print_dry_run(parsed)
        if errors:
            print("\n[ERROR]", file=sys.stderr)
            for e in errors:
                print(f"- {e}", file=sys.stderr)
        return 1 if errors else 0

    validate_db_env()
    print(f"[DB] {masked_db_target()}")
    conn = connect_db()
    try:
        ensure_schema(conn)
        success = 0
        for parsed in parsed_docs:
            try:
                qid, action = upsert_questionnaire(conn, parsed)
                tag_names = [t["tag_name"] for t in parsed.tags]
                print(
                    f"[OK] seq={parsed.answers.get('seq_name')} id={qid} "
                    f"tags={len(tag_names)} file={parsed.source_file}"
                )
                if tag_names:
                    print("     " + ", ".join(tag_names))
                success += 1
            except Exception as exc:
                conn.rollback()
                errors.append(f"{parsed.source_file}: DB 적재 실패 - {exc}")

        print(f"\n완료: 성공 {success}건 / 실패 {len(errors)}건")
        if errors:
            for e in errors:
                print(f"[ERROR] {e}", file=sys.stderr)
        return 1 if errors else 0
    finally:
        conn.close()


if __name__ == "__main__":
    raise SystemExit(main())
